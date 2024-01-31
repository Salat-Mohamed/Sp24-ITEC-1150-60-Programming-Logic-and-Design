import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Fetch the country list from the API (same as in the previous code)
response = requests.get("https://country-list-1150.azurewebsites.net/api/country")
data = response.json()

# Create a new Word document
document = Document()

# Add a title
document.add_heading("Countries and their Capital Cities", level=0)

# Iterate through the country list and add formatted text
for country in data:
    try:
        country_name = country["name"]
        capital_city = country["capital"]

        # Create a paragraph for each country-capital pair
        paragraph = document.add_paragraph(f"The capital of {country_name} is {capital_city}.")

        # Apply formatting (optional)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except KeyError:
        print(f"Missing data for country: {country}")

# Save the document
document.save("countries_and_capitals.docx")

print("Word document created successfully!")
